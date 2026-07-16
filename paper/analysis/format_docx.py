#!/usr/bin/env python3
"""Apply submission-oriented formatting to Pandoc-generated NWKIT documents."""

from __future__ import annotations

import argparse
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MAIN_ALT_TEXT = [
    (
        "NWKIT architecture: Newick trees and metadata pass through shared input rules, "
        "five shaded command families, and shared output rules to Newick, tables, and figures."
    ),
    (
        "PEPC C4-state workflow: branches of a 71-tip tree vary along a marginal C4-probability "
        "scale; C3 circles and C4 triangles mark observed states, and eight selected C4 tips are "
        "outlined and labelled by species and accession."
    ),
    (
        "Evaluation summary: agreement bars show complete agreement in 191 checks, all 12 "
        "input-case outcomes match expectations, and line charts show preflight and consensus "
        "timing across input sizes."
    ),
]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, margin_twips: int = 55) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil" if edge in {"left", "right", "insideV"} else "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "808080")


def add_continuous_line_numbers(section) -> None:
    sect_pr = section._sectPr
    for previous in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(previous)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    line_numbers.set(qn("w:distance"), "360")
    sect_pr.append(line_numbers)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = paragraph._p.get_or_add_pPr()
    suppress = OxmlElement("w:suppressLineNumbers")
    p_pr.append(suppress)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def set_run_font(run, size: float, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def configure_styles(document: Document, main: bool) -> None:
    base_size = 12 if main else 8.5
    line_spacing = 1.5 if main else 1.0
    for style_name in ("Normal", "Body Text", "First Paragraph"):
        if style_name not in document.styles:
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(base_size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_after = Pt(0 if main else 3)
        style.paragraph_format.widow_control = True
        style.paragraph_format.first_line_indent = Inches(0.5) if main else None

    title = document.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(14 if main else 13)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    if "Hyperlink" in document.styles:
        hyperlink = document.styles["Hyperlink"]
        hyperlink.font.name = "Times New Roman"
        hyperlink._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        hyperlink.font.color.rgb = RGBColor(0, 0, 0)
        hyperlink.font.underline = False

    for level in (1, 2, 3):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12 if main else (10 if level == 1 else 9))
        style.font.bold = not (main and level == 3)
        style.font.italic = bool(main and level == 3)
        style.font.small_caps = bool(main and level == 2)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(9 if main else 6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0


def format_main(document: Document) -> None:
    configure_styles(document, main=True)
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    add_continuous_line_numbers(section)
    add_page_number(section)

    for paragraph in list(document.paragraphs):
        if paragraph.style.name == "Image Caption":
            paragraph._element.getparent().remove(paragraph._element)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.first_line_indent = None
        elif paragraph.text in {"Kenji Fukushima", "Corresponding author: kfuku52@gmail.com"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
        elif paragraph.text.startswith(("Author running head:", "Title running head:")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            if paragraph.text.startswith("Title running head:"):
                paragraph.paragraph_format.space_after = Pt(12)
        elif paragraph.text.startswith("Keywords:"):
            paragraph.paragraph_format.first_line_indent = None
        elif paragraph.text.startswith(("Figure ", "Table 1.")):
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_together = True
            if paragraph.text.startswith("Figure "):
                paragraph.paragraph_format.keep_with_next = True
        elif paragraph.text.startswith("Alt text:"):
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.keep_together = True
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, 12)

    max_width = Inches(6.35)
    max_height = Inches(6.1)
    for index, shape in enumerate(document.inline_shapes):
        scale = min(max_width / shape.width, max_height / shape.height)
        if scale < 1:
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)
        doc_pr_nodes = shape._inline.xpath(".//wp:docPr")
        if doc_pr_nodes and index < len(MAIN_ALT_TEXT):
            doc_pr_nodes[0].set("descr", MAIN_ALT_TEXT[index])
            doc_pr_nodes[0].set("title", f"Figure {index + 1}")

    table = document.tables[0]
    table.style = "Table"
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.55, 0.92, 0.92, 0.92, 1.04]
    for col_index, width in enumerate(widths):
        table.columns[col_index].width = Inches(width)
        table._tbl.tblGrid.gridCol_lst[col_index].w = Inches(width)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 45)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    set_run_font(run, 8.2, bold=True if row_index == 0 else None)


def column_widths(table, usable_inches: float) -> list[float]:
    lengths = []
    for col_index in range(len(table.columns)):
        max_len = max(len(row.cells[col_index].text) for row in table.rows)
        lengths.append(max(5.0, min(42.0, math.sqrt(max_len + 1) * 3.0)))
    total = sum(lengths)
    return [usable_inches * weight / total for weight in lengths]


def format_supplement(document: Document) -> None:
    configure_styles(document, main=False)
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    add_page_number(section)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.style.name == "Heading 2":
            paragraph.paragraph_format.page_break_before = True
        for run in paragraph.runs:
            size = 13 if paragraph.style.name == "Title" else 8.5
            set_run_font(run, size)

    for table in document.tables:
        table.style = "Table"
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = column_widths(table, 9.9)
        set_repeat_table_header(table.rows[0])
        font_size = 6.5 if len(table.columns) >= 10 else 7.2
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            for col_index, cell in enumerate(row.cells):
                cell.width = Inches(widths[col_index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_margins(cell, 30)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.keep_together = True
                    for run in paragraph.runs:
                        set_run_font(run, font_size, bold=True if row_index == 0 else None)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--supplement", action="store_true")
    args = parser.parse_args()

    document = Document(args.input)
    if args.supplement:
        format_supplement(document)
    else:
        format_main(document)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
